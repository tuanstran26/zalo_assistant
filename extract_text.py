import os
import fitz
import re
import tempfile
import json
import datetime
from docx import Document
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from storing_supabase import process_document_and_store
from excel_extractor import extract_excel
from excel_chunking import store_excel_file

# Cấu hình
DOWNLOAD_DIR = "downloads"
FILE_EXTENSIONS = [".pdf", ".docx", ".txt", ".jpg", ".png", ".jpeg"]
POPPLER_PATH = r"poppler-24.08.0\Library\bin"  # Windows cần khai báo
current_dir = os.path.dirname(os.path.abspath(__file__))

tesseract_path = os.path.join(current_dir, 'Tesseract-OCR', 'tesseract.exe')
pytesseract.pytesseract.tesseract_cmd = tesseract_path

def clean_text(raw_text):
    cleaned_text = re.sub(r'^\s*(\.\s*){3,}\s*$', '', raw_text, flags=re.MULTILINE)
    cleaned_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', cleaned_text.strip())
    return cleaned_text

def preprocess_text(text):
    text = re.sub(r'(?:Page|Trang)?\s*-?\s*\d+\s*-?', '', text, flags=re.IGNORECASE)
    text = re.sub(r"[^\w\s.,!?%\-–()]", "", text)
    text = re.sub(r'\n{2,}', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return preprocess_text(clean_text(f.read()))

def extract_text_from_image_files(file_path):
    text = ""
    if file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='vie+eng')
        image.close()
    elif file_path.lower().endswith('.pdf'):
        images = convert_from_path(file_path, dpi=150, grayscale=True, poppler_path=POPPLER_PATH)
        for image in images:
            temp_path = tempfile.mktemp(suffix='.png')
            image.save(temp_path, 'PNG')
            text += pytesseract.image_to_string(Image.open(temp_path), lang='vie+eng') + "\n"
            os.unlink(temp_path)
            image.close()
    elif file_path.lower().endswith('.docx'):
        doc = Document(file_path)
        for rel in doc.part._rels:
            target = doc.part._rels[rel].target_ref
            if "image" in target:
                img_bytes = doc.part.related_parts[target].blob
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                    temp_img.write(img_bytes)
                    temp_img_path = temp_img.name
                text += pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng') + "\n"
                os.unlink(temp_img_path)
    else:
        raise ValueError(f"Unsupported file type for image OCR: {file_path}")
    return preprocess_text(clean_text(text))

def extract_text_from_docx_with_image_and_text(file_path):
    if not file_path.lower().endswith('.docx'):
        raise ValueError("File không phải định dạng .docx")
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
 

    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        target = rel_obj.target_ref
        if "image" in target:
            try:
                img_part = rel_obj.target_part
                img_bytes = img_part.blob
            except Exception as e:
                print(f"[Warning] Không thể đọc ảnh {target}: {e}")
                continue

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_img.write(img_bytes)
                temp_img_path = temp_img.name

            ocr_result = pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng')
            if ocr_result.strip():
                full_text.append(ocr_result.strip())
            os.unlink(temp_img_path)


    combined_text = "\n".join(full_text)
    return preprocess_text(clean_text(combined_text))

def extract_text_from_pdf_with_image_and_text(file_path):
    if not file_path.lower().endswith('.pdf'):
        raise ValueError("File không phải định dạng PDF")
    doc = fitz.open(file_path)
    combined_text = ""
    for page in doc:
        text = page.get_text().strip()
        if text:
            combined_text += text + "\n"
        else:
            pix = page.get_pixmap(dpi=150)
            temp_img_path = tempfile.mktemp(suffix=".png")
            pix.save(temp_img_path)
            ocr_result = pytesseract.image_to_string(Image.open(temp_img_path), lang='vie+eng')
            combined_text += ocr_result.strip() + "\n"
            os.unlink(temp_img_path)
    doc.close()
    return preprocess_text(clean_text(combined_text))

def extract_text_from_text_only_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return preprocess_text(clean_text(text))

def extract_text_from_text_only_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return preprocess_text(clean_text(text))



def extract_text_from_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image, lang='vie+eng')
    image.close()
    return preprocess_text(clean_text(text))


def get_file_metadata(file_path):
    stats = os.stat(file_path)
    metadata = {
        "filename": os.path.basename(file_path),
        "size_bytes": stats.st_size,
        "create_date": datetime.datetime.fromtimestamp(stats.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
        "modified_date": datetime.datetime.fromtimestamp(stats.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        "extension": os.path.splitext(file_path)[1].lower()
    }
    return metadata



def check_pdf_content_type(file_path):
    doc = fitz.open(file_path)
    has_text = False
    has_image = False
    for page in doc:
        if page.get_text().strip():
            has_text = True
        if page.get_images(full=True):
            has_image = True
        if has_text and has_image:
            break
    doc.close()
    if has_text and has_image:
        return 1
    elif has_text:
        return 2
    elif has_image:
        return 3
    else:
        return 0

def check_docx_content_type(file_path):
    doc = Document(file_path)
    has_text = False
    has_image = False
    for para in doc.paragraphs:
        if para.text.strip():
            has_text = True
            break
    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            has_image = True
            break
    if has_text and has_image:
        return 1
    elif has_text:
        return 2
    elif has_image:
        return 3
    else:
        return 0

def extract_text_with_metadata(file_name: str):
    file_path = os.path.join(DOWNLOAD_DIR, file_name)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_path} không tồn tại")

    lower_path = file_path.lower()
    if lower_path.endswith(('.jpg', '.png', '.jpeg')):
        text = extract_text_from_image_files(file_path)
    elif lower_path.endswith('.txt'):
        text = extract_text_from_txt(file_path)
    elif lower_path.endswith('.pdf'):
        file_type = check_pdf_content_type(file_path)
        if file_type == 1:
            text = extract_text_from_pdf_with_image_and_text(file_path)
        elif file_type == 2:
            text = extract_text_from_text_only_pdf(file_path)
        elif file_type == 3:
            text = extract_text_from_image_files(file_path)
        else:
            text = ""
    elif lower_path.endswith('.docx'):
        file_type = check_docx_content_type(file_path)
        if file_type == 1:
            text = extract_text_from_docx_with_image_and_text(file_path)
        elif file_type == 2:
            text = extract_text_from_text_only_docx(file_path)
        elif file_type == 3:
            text = extract_text_from_image_files(file_path)
        else:
            text = ""
    elif lower_path.endswith('.xlsx'):
            text = extract_excel(file_name)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")


    metadata = get_file_metadata(file_path)

    doc ={
        "file_name": file_name,
        "metadata": metadata,
        "content": text
    }
    if lower_path.endswith('.xlsx'):
        store_excel_file(file_name, text, metadata)
    else:
        process_document_and_store(doc)
    return doc


if __name__ == "__main__":
    test_file = "DÀN Ý VĂN 6, CẢ 2 KỲ.docx"
    result = extract_text_with_metadata(test_file)

    print(json.dumps(result, indent=2, ensure_ascii=False)[:2000]) 
